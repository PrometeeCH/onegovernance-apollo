import json
import os
import uuid
from typing import Any, List

import fitz  # PyMuPDF for extracting PDF text and metadata
import pandas as pd
import streamlit as st
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes.models import (
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SimpleField,
)
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain.retrievers.multi_vector import MultiVectorRetriever
from langchain.storage import InMemoryByteStore
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores.azuresearch import AzureSearch
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_text_splitters import (
    CharacterTextSplitter,
    RecursiveCharacterTextSplitter,
)

load_dotenv()


class VectorStore:
    def __init__(self) -> None:
        embeddings: AzureOpenAIEmbeddings = AzureOpenAIEmbeddings(
            azure_deployment=os.getenv("AZURE_DEPLOYMENT"),
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
            azure_endpoint=os.getenv("AZURE_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        )
        embedding_function = embeddings.embed_query
        _DEFAULT = "one-test-3"
        index_name: str = os.getenv("INDEX_NAME", _DEFAULT)
        fields = [
            SimpleField(
                name="id",
                type=SearchFieldDataType.String,
                key=True,
                filterable=True,
            ),
            SearchableField(
                name="content",
                type=SearchFieldDataType.String,
                searchable=True,
            ),
            SearchField(
                name="content_vector",
                type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                searchable=True,
                vector_search_dimensions=len(embedding_function("Text")),
                vector_search_profile_name="myHnswProfile",
            ),
            SearchableField(
                name="metadata",
                type=SearchFieldDataType.String,
                searchable=True,
            ),
        ]
        vector_store: AzureSearch = AzureSearch(
            azure_search_endpoint=os.getenv("VECTOR_STORE_ADDRESS"),
            azure_search_key=os.getenv("VECTOR_STORE_PASSWORD"),
            index_name=index_name,
            embedding_function=embeddings.embed_query,
            fields=fields,  # This allows us to put any schema we want with Langchain
        )
        search_client: SearchClient = SearchClient(
            endpoint=os.getenv("AZURE_SEARCH_ENDPOINT"),
            index_name=index_name,
            credential=AzureKeyCredential(os.getenv("AZURE_SEARCH_KEY")),
        )
        self._search_client = search_client
        self._vector_store = vector_store
        self._embedding_function = embedding_function
        self._index_name = index_name
        self._fields = fields

    def get_vector_store(self) -> AzureSearch:
        return self._vector_store

    def get_embedding_function(self) -> Any:
        return self._embedding_function

    def get_index_name(self) -> str:
        return self._index_name

    def get_search_client(self) -> SearchClient:
        return self._search_client

    def get_fields(self) -> List:
        return self._fields

    @staticmethod
    def extract_title(pdf_path: str) -> str:
        title: str = ""
        with fitz.open(pdf_path) as doc:
            title = doc.metadata["title"]  # Attempt to use metadata title
            if not title:  # Fallback to the first line of the first page
                first_page = doc[0]
                title = first_page.get_text().split("\n")[0]
        return title

    @staticmethod
    def extract_title_from_csv(csv_path: str) -> str:
        with open(csv_path, "r") as f:
            title = f.readline().split(",")[0].strip()
        return title

    @staticmethod
    def extract_title_from_xlsx(xlsx_path: str) -> str:
        title = xlsx_path
        return title

    @staticmethod
    def extract_title_from_docx(docx_path: str) -> str:
        doc = DocxDocument(docx_path)

        title = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Si le paragraphe contient du texte
                title = paragraph.text  # Alors cela devient le titre
                break  # Sortir de la boucle car nous avons trouvé le titre

        return title

    @staticmethod
    def load_document(file_path: str) -> List[str]:
        file_type = file_path.split(".")[-1].lower()
        documents = []

        if file_type == "pdf":
            loader = PyPDFLoader(file_path=file_path)
            documents = loader.load()
        elif file_type == "txt":
            loader = TextLoader(file_path)
            documents = loader.load()
        elif file_type == "csv":
            documents = load_csv_as_documents(file_path)
        elif file_type == "xlsx":
            documents = load_xlsx_as_documents(file_path)
        elif file_type == "docx":
            documents = load_xlsx_as_documents(file_path)
        return documents

    @staticmethod
    def add_metadata(docs: Any, title: str, start_index: int) -> List[str]:
        anonymized_docs = []

        for i, doc in enumerate(docs):
            # Anonymize doc
            # anonymized_doc = Anonymizer(
            #    original_text=doc.page_content
            # ).get_anonymized_text()

            # doc = doc.copy(update={"page_content": anonymized_doc})
            doc.metadata.update(
                {"title": f"{title}-{start_index + i + 1}", "source": title}
            )  # Ensure different titles for multiple chunks of the same document.
            anonymized_docs.append(doc)

        return anonymized_docs

    def char_text_splitter(self, docs: Any, title: str, start_index: int) -> List[str]:
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        sub_docs = text_splitter.split_documents(docs)
        return self.add_metadata(sub_docs, title, start_index=start_index)

    def rec_char_text_splitter(
        self, docs: Any, title: str, start_index: int
    ) -> List[str]:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=400)
        sub_docs = text_splitter.split_documents(docs)
        return self.add_metadata(sub_docs, title, start_index=start_index)

    def summary_chunk(self, docs: Any, title: str, start_index: int) -> List[str]:
        chain = (
            {"doc": lambda x: x.page_content}
            | ChatPromptTemplate.from_template(
                "Summarize the following document:\n\n{doc}"
            )
            | AzureChatOpenAI(
                azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT_CHAT"),
                openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
                azure_deployment=os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"),
                openai_api_key=os.getenv("AZURE_OPENAI_API_KEY_CHAT"),
            )
            | StrOutputParser()
        )
        summaries = chain.batch(docs, {"max_concurrency": 5})
        summary_docs = [Document(page_content=s) for i, s in enumerate(summaries)]
        return self.add_metadata(summary_docs, title, start_index=start_index)

    def push_document(self, document_path: str, title: str) -> None:
        # Load documents depending on type (PDF or TXT)
        docs = self.load_document(document_path)
        id_key = "doc_id"
        retriever = MultiVectorRetriever(
            vectorstore=self._vector_store,
            id_key=id_key,
            byte_store=InMemoryByteStore(),
        )

        progress_bar = st.progress(0)
        # Char Splitter
        anonymized_docs1 = self.char_text_splitter(docs, title, start_index=0)
        retriever.vectorstore.add_documents(documents=anonymized_docs1)
        retriever.docstore.mset(list(zip(str(uuid.uuid4()), docs)))
        progress_bar.progress(33)
        # Recursive Char Splitter
        anonymized_docs2 = self.rec_char_text_splitter(
            docs, title, start_index=len(anonymized_docs1)
        )
        retriever.vectorstore.add_documents(documents=anonymized_docs2)
        retriever.docstore.mset(list(zip(str(uuid.uuid4()), docs)))
        progress_bar.progress(66)
        # Summary
        anonymized_docs3 = self.summary_chunk(
            docs, title, start_index=len(anonymized_docs1) + len(anonymized_docs2)
        )
        retriever.vectorstore.add_documents(documents=anonymized_docs3)
        retriever.docstore.mset(list(zip(str(uuid.uuid4()), docs)))
        progress_bar.progress(100)

        progress_bar.empty()

    def fetch_documents(self, max_results: int = 10) -> List:
        results = []
        results_title = []  # To ensure we don't have duplicates
        search_results = self._search_client.search(search_text="*", top=max_results)
        for result in search_results:
            metadata = json.loads(result["metadata"])
            if metadata["title"] not in results_title:
                results.append(result)
                results_title.append(metadata["title"])

        return results

    def delete_document(self, document_key: str) -> str:
        batch = [{"id": document_key}]
        try:
            self._search_client.delete_documents(documents=batch)
            print(f"Document ID {document_key} deleted.")
            return ""
        except Exception as e:
            return f"Error detected while deleting document {document_key}: {e}"


def load_csv_as_documents(file_path: str) -> List[Document]:
    # Charger le fichier CSV
    df = pd.read_csv(file_path)

    # Créer une liste pour stocker les objets Document
    documents = []

    # Itérer sur chaque ligne du DataFrame
    for index, row in df.iterrows():
        # Concaténer toutes les colonnes pour former le contenu de la page
        page_content = " ".join(row.astype(str))

        # Créer un objet Document de langchain_core avec le contenu de la page et le numéro de ligne comme métadonnée
        documents.append(
            Document(page_content=page_content, metadata={"line_number": index})
        )

    return documents


def load_docx_as_documents(file_path: str) -> List[Document]:
    # Load the docx file
    doc = DocxDocument(file_path)

    # Create a list to store Document objects
    documents = []

    # Iterate over the paragraphs in the doc
    for index, para in enumerate(doc.paragraphs):
        # Use each paragraph as the page content
        page_content = para.text

        # Create a Document object from langchain_core with the page content and paragraph number as metadata
        documents.append(
            Document(page_content=page_content, metadata={"paragraph_number": index})
        )

    return documents


def load_xlsx_as_documents(file_path: str) -> List[Document]:
    # Load the xlsx file
    df = pd.read_excel(file_path)

    # Create a list to store Document objects
    documents = []

    # Iterate over each row in the dataframe
    for index, row in df.iterrows():
        # Concatenate all columns to form the page content
        page_content = " ".join(row.astype(str))

        # Create a Document object from langchain_core with the page content and row number as metadata
        documents.append(
            Document(page_content=page_content, metadata={"line_number": index})
        )

    return documents
